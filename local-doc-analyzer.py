import os
import json
import hashlib
import pickle
from typing import List, Dict, Optional, Union
from dataclasses import dataclass, asdict
import numpy as np
from sentence_transformers import SentenceTransformer
from transformers import pipeline
from pypdf import PdfReader
import docx
import faiss
import torch
from tqdm import tqdm
from pathlib import Path
import pandas as pd
import openpyxl
import markdown
from bs4 import BeautifulSoup
import re
from flask import Flask, request, jsonify, render_template_string
import threading

@dataclass
class Document:
    content: str
    metadata: Dict
    embeddings: Optional[np.ndarray] = None
    summary: Optional[str] = None
    
    def to_dict(self):
        """Convert Document to dictionary for serialization."""
        return {
            'content': self.content,
            'metadata': self.metadata,
            'summary': self.summary,
            'embeddings': self.embeddings.tobytes() if self.embeddings is not None else None
        }
    
    @classmethod
    def from_dict(cls, data):
        """Create Document from dictionary."""
        if data['embeddings'] is not None:
            data['embeddings'] = np.frombuffer(data['embeddings'], dtype=np.float32)
        return cls(**data)

class DocumentPreprocessor:
    """Advanced document preprocessing capabilities."""
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize text content."""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s.,!?-]', '', text)
        return text.strip()
    
    @staticmethod
    def extract_structured_content(text: str) -> Dict:
        """Extract structured information from text."""
        # Extract potential headers (capitalized lines)
        headers = re.findall(r'^[A-Z][A-Z\s]+(?=:|\n)', text, re.MULTILINE)
        
        # Extract potential lists (bullet points or numbered items)
        lists = re.findall(r'(?:^\d+\.|\*)\s+.+', text, re.MULTILINE)
        
        return {
            'headers': headers,
            'lists': lists,
            'full_text': text
        }

class LocalDocumentAnalyzer:
    def __init__(
        self,
        embedding_model: str = "all-MiniLM-L6-v2",
        cache_dir: str = ".cache",
        device: str = None
    ):
        """Initialize the document analyzer with specified models and caching."""
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(exist_ok=True)
        
        self.device = device or ('cuda' if torch.cuda.is_available() else 'cpu')
        
        self.embedding_model = SentenceTransformer(embedding_model)
        self.qa_pipeline = pipeline(
            "question-answering",
            model="deepset/roberta-base-squad2",
            device=0 if self.device == 'cuda' else -1
        )
        self.summarization_pipeline = pipeline(
            "summarization",
            model="facebook/bart-large-cnn",
            device=0 if self.device == 'cuda' else -1
        )
        
        self.documents = []
        self.index = None
        self.preprocessor = DocumentPreprocessor()
    
    def _get_cache_path(self, file_path: str) -> Path:
        """Generate cache file path based on input file hash."""
        file_hash = hashlib.md5(file_path.encode()).hexdigest()
        return self.cache_dir / f"{file_hash}.pickle"
    
    def _load_from_cache(self, file_path: str) -> Optional[Document]:
        """Load document from cache if available."""
        cache_path = self._get_cache_path(file_path)
        if cache_path.exists():
            with open(cache_path, 'rb') as f:
                return Document.from_dict(pickle.load(f))
        return None
    
    def _save_to_cache(self, file_path: str, document: Document):
        """Save document to cache."""
        cache_path = self._get_cache_path(file_path)
        with open(cache_path, 'wb') as f:
            pickle.dump(document.to_dict(), f)
    
    def load_document(self, file_path: str) -> Document:
        """Load document from various file formats with caching."""
        # Check cache first
        cached_doc = self._load_from_cache(file_path)
        if cached_doc:
            return cached_doc
        
        file_ext = os.path.splitext(file_path)[1].lower()
        content = ""
        
        try:
            if file_ext == '.pdf':
                reader = PdfReader(file_path)
                for page in reader.pages:
                    content += page.extract_text() + "\n"
            elif file_ext == '.docx':
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    content += para.text + "\n"
            elif file_ext in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path)
                content = df.to_string()
            elif file_ext == '.csv':
                df = pd.read_csv(file_path)
                content = df.to_string()
            elif file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif file_ext in ['.html', '.htm']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    soup = BeautifulSoup(f.read(), 'html.parser')
                    content = soup.get_text()
            elif file_ext == '.md':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = markdown.markdown(f.read())
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
        
        except Exception as e:
            raise Exception(f"Error processing {file_path}: {str(e)}")
        
        # Preprocess content
        content = self.preprocessor.clean_text(content)
        structured_content = self.preprocessor.extract_structured_content(content)
        
        # Create document
        doc = Document(
            content=content,
            metadata={
                "file_path": file_path,
                "file_type": file_ext,
                "structured_content": structured_content
            }
        )
        
        # Generate summary
        doc.summary = self.summarize_text(content)
        
        # Cache document
        self._save_to_cache(file_path, doc)
        
        return doc
    
    def summarize_text(self, text: str, max_length: int = 150) -> str:
        """Generate a summary of the text."""
        chunks = [text[i:i + 1024] for i in range(0, len(text), 1024)]
        summaries = []
        
        for chunk in chunks:
            summary = self.summarization_pipeline(chunk, max_length=max_length, min_length=30)[0]['summary_text']
            summaries.append(summary)
        
        return " ".join(summaries)
    
    def process_documents(self, file_paths: List[str], chunk_size: int = 512):
        """Process multiple documents and build the FAISS index."""
        print("Loading and processing documents...")
        for file_path in tqdm(file_paths):
            doc = self.load_document(file_path)
            # Split content into chunks
            chunks = [doc.content[i:i + chunk_size] 
                     for i in range(0, len(doc.content), chunk_size)]
            
            for chunk in chunks:
                chunk_doc = Document(
                    content=chunk,
                    metadata=doc.metadata.copy(),
                    summary=doc.summary
                )
                self.documents.append(chunk_doc)
        
        print("Computing embeddings...")
        embeddings = []
        for doc in tqdm(self.documents):
            embedding = self.embedding_model.encode(doc.content)
            doc.embeddings = embedding
            embeddings.append(embedding)
            
        embeddings_array = np.array(embeddings).astype('float32')
        
        # Build FAISS index
        print("Building FAISS index...")
        self.index = faiss.IndexFlatL2(embeddings_array.shape[1])
        self.index.add(embeddings_array)
    
    def search_documents(self, query: str, k: int = 3) -> List[Document]:
        """Search for relevant documents using the query."""
        query_embedding = self.embedding_model.encode([query])[0].astype('float32')
        distances, indices = self.index.search(query_embedding.reshape(1, -1), k)
        
        return [self.documents[i] for i in indices[0]]
    
    def answer_question(self, query: str, k: int = 3) -> Dict:
        """Answer a question using relevant document chunks."""
        relevant_docs = self.search_documents(query, k)
        context = " ".join([doc.content for doc in relevant_docs])
        
        answer = self.qa_pipeline(
            question=query,
            context=context,
            max_answer_length=100
        )
        
        return {
            "answer": answer["answer"],
            "confidence": answer["score"],
            "source_documents": [doc.metadata for doc in relevant_docs]
        }

# Web Interface
app = Flask(__name__)
analyzer = None

HTML_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <title>Document Analyzer</title>
    <style>
        body { font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }
        .container { margin-top: 20px; }
        .result { margin-top: 20px; padding: 10px; border: 1px solid #ddd; }
        .file-list { margin-top: 10px; }
        textarea { width: 100%; height: 100px; }
        button { margin-top: 10px; }
    </style>
</head>
<body>
    <h1>Document Analyzer</h1>
    
    <div class="container">
        <h2>Upload Documents</h2>
        <form action="/upload" method="post" enctype="multipart/form-data">
            <input type="file" name="files" multiple>
            <button type="submit">Upload & Process</button>
        </form>
    </div>

    <div class="container">
        <h2>Ask a Question</h2>
        <textarea id="question" placeholder="Enter your question here..."></textarea>
        <button onclick="askQuestion()">Ask</button>
        <div id="answer" class="result"></div>
    </div>

    <div class="container">
        <h2>Document Summaries</h2>
        <div id="summaries" class="result"></div>
    </div>

    <script>
        async function askQuestion() {
            const question = document.getElementById('question').value;
            const response = await fetch('/ask', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify({question: question})
            });
            const result = await response.json();
            document.getElementById('answer').innerHTML = `
                <p><strong>Answer:</strong> ${result.answer}</p>
                <p><strong>Confidence:</strong> ${(result.confidence * 100).toFixed(2)}%</p>
                <p><strong>Sources:</strong></p>
                <ul>
                    ${result.source_documents.map(doc => `<li>${doc.file_path}</li>`).join('')}
                </ul>
            `;
        }
    </script>
</body>
</html>
"""

@app.route('/')
def home():
    return render_template_string(HTML_TEMPLATE)

@app.route('/upload', methods=['POST'])
def upload_files():
    if 'files' not in request.files:
        return jsonify({'error': 'No files provided'}), 400
    
    files = request.files.getlist('files')
    file_paths = []
    
    for file in files:
        file_path = os.path.join('uploads', file.filename)
        os.makedirs('uploads', exist_ok=True)
        file.save(file_path)
        file_paths.append(file_path)
    
    analyzer.process_documents(file_paths)
    return jsonify({'message': 'Files processed successfully'})

@app.route('/ask', methods=['POST'])
def ask_question():
    question = request.json.get('question')
    if not question:
        return jsonify({'error': 'No question provided'}), 400
    
    result = analyzer.answer_question(question)
    return jsonify(result)

def start_server():
    """Start the Flask server."""
    global analyzer
    analyzer = LocalDocumentAnalyzer()
    app.run(host='0.0.0.0', port=5000)

if __name__ == "__main__":
    start_server()
